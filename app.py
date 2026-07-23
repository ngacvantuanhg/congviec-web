"""
PHẦN MỀM QUẢN LÝ CÔNG TÁC CÁ NHÂN — BẢN WEB (Streamlit + Supabase)
Ban Tuyên giáo và Dân vận Tỉnh ủy Tuyên Quang
"""
import io
import calendar
import datetime
from collections import Counter

import streamlit as st
import pandas as pd
import plotly.express as px
from supabase import create_client
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════
#  CẤU HÌNH CHUNG
# ══════════════════════════════════════════════════════
st.set_page_config(page_title="Quản lý Công tác Cá nhân — Ban TG&DV Tuyên Quang",
                    page_icon="⭐", layout="wide")

NAVY  = "#1B3A6B"
RED   = "#C8102E"
GOLD  = "#F5A623"
GREEN = "#1A9651"
ORANGE= "#E67E22"
DANGER= "#C0392B"

LOAI_CV = [
    "Tham mưu văn bản", "Xây dựng ứng dụng chuyển đổi số", "Họp Ban",
    "Họp Chi bộ", "Họp chuyên môn", "Công tác trong tỉnh",
    "Công tác ngoài tỉnh", "Dự tập huấn", "Giao dịch với đơn vị ngoài cơ quan",
    "Khác",
]
TRANG_THAI = ["Chưa thực hiện", "Đang thực hiện", "Hoàn thành", "Trễ hạn"]
UU_TIEN = ["Cao", "Bình thường", "Thấp"]
VN_THU = ["Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu", "Thứ Bảy", "Chủ Nhật"]
LOAI_VIEC_BC = ["Kế hoạch/thường xuyên", "Phát sinh, đột xuất"]  # dùng cho Biểu 01 tự đánh giá

st.markdown(f"""
<style>
.stApp {{ background:#F2F5F9; }}
[data-testid="stSidebar"] {{ background:{NAVY}; }}
[data-testid="stSidebar"] * {{ color:#E4ECF7 !important; }}
h1,h2,h3 {{ color:{NAVY}; }}
div.stButton>button {{ background:{NAVY}; color:white; border-radius:8px; border:none; }}
div.stButton>button:hover {{ background:{RED}; color:white; }}
[data-testid="stMetricValue"] {{ color:{NAVY}; }}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
#  KẾT NỐI SUPABASE
# ══════════════════════════════════════════════════════
@st.cache_resource
def get_client():
    url = st.secrets["supabase"]["url"]
    key = st.secrets["supabase"]["key"]
    return create_client(url, key)

sb = get_client()

# ── Xác thực đơn giản bằng mật khẩu chung (tùy chọn) ──
if "app_password" in st.secrets.get("auth", {}):
    if not st.session_state.get("authed"):
        st.title("🔒 Đăng nhập")
        with st.form("login_form"):
            pw = st.text_input("Mật khẩu", type="password")
            submitted = st.form_submit_button("Đăng nhập")
        if submitted:
            if pw == st.secrets["auth"]["app_password"]:
                st.session_state["authed"] = True
                st.rerun()
            else:
                st.error("Sai mật khẩu")
        st.stop()

# ══════════════════════════════════════════════════════
#  TRUY VẤN DỮ LIỆU
# ══════════════════════════════════════════════════════
def cfg_get(key, default=""):
    r = sb.table("config").select("value").eq("key", key).execute().data
    return r[0]["value"] if r else default

def cfg_set(key, val):
    sb.table("config").upsert({"key": key, "value": val}).execute()

def db_query(tu=None, den=None):
    q = sb.table("tasks").select("*")
    if tu:
        q = q.gte("ngay_kt", str(tu))
    if den:
        q = q.lte("ngay_bd", str(den))
    data = q.execute().data
    data.sort(key=lambda t: (t["ngay_bd"], t.get("gio_bd") or ""))
    return data

def db_all():
    data = sb.table("tasks").select("*").execute().data
    data.sort(key=lambda t: (t["ngay_bd"], t.get("gio_bd") or ""))
    return data

def db_them(data):
    return sb.table("tasks").insert(data).execute().data

def db_sua(tid, data):
    sb.table("tasks").update(data).eq("id", tid).execute()

def db_xoa(tid):
    sb.table("tasks").delete().eq("id", tid).execute()

def nwt_query(tuan, nam):
    r = sb.table("next_week_tasks").select("*").eq("tuan", tuan).eq("nam", nam).execute().data
    r.sort(key=lambda x: x["id"])
    return r

def nwt_them(data):
    sb.table("next_week_tasks").insert(data).execute()

def nwt_sua(nid, data):
    sb.table("next_week_tasks").update(data).eq("id", nid).execute()

def nwt_xoa(nid):
    sb.table("next_week_tasks").delete().eq("id", nid).execute()

# ══════════════════════════════════════════════════════
#  TIỆN ÍCH WORD
# ══════════════════════════════════════════════════════
def _shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def _style_header_row(row, fill="1B3A6B"):
    for cell in row.cells:
        _shade_cell(cell, fill)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def _fill_row(cells, values, size=10, center_cols=()):
    for i, (cell, v) in enumerate(zip(cells, values)):
        cell.text = ""
        p = cell.paragraphs[0]
        if i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(v) if v is not None else "")
        r.font.size = Pt(size)

def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)

def _fmt_d(s):
    if not s:
        return ""
    try:
        return datetime.date.fromisoformat(str(s)).strftime("%d/%m/%Y")
    except Exception:
        return str(s)

# ══════════════════════════════════════════════════════
#  BÁO CÁO THEO MẪU "TỔNG HỢP BC CỦA CBCC VP" (mới)
# ══════════════════════════════════════════════════════
def build_cbcc_report(ten, tasks_ky, nwt_rows, tuan_so, nam):
    """Tạo file .docx đúng theo bố cục Mẫu - Tổng hợp BC của CBCC VP."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2); sec.right_margin = Cm(2)

    p = doc.add_paragraph()
    r = p.add_run("Họ và tên: ")
    r.bold = True; r.font.size = Pt(12)
    r2 = p.add_run(ten)
    r2.font.size = Pt(12)

    # ── 1. Kết quả thực hiện nhiệm vụ ──
    h1 = doc.add_paragraph()
    r = h1.add_run("1. Kết quả thực hiện nhiệm vụ")
    r.bold = True; r.font.size = Pt(12)

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    heads = ["STT", "Nội dung", "Lãnh đạo giao", "Thời gian được giao",
             "Thời hạn của VB", "Thời gian hoàn thành", "Lãnh đạo thẩm định, duyệt"]
    _fill_row(tbl.rows[0].cells, heads, size=10)
    _style_header_row(tbl.rows[0])
    for i, t in enumerate(tasks_ky, 1):
        cells = tbl.add_row().cells
        _fill_row(cells, [
            str(i),
            t["title"],
            t.get("lanh_dao_giao") or "",
            _fmt_d(t.get("ngay_bd")),
            _fmt_d(t.get("thoi_han_vb")),
            _fmt_d(t.get("ngay_kt")) if t.get("trang_thai") == "Hoàn thành" else "",
            t.get("lanh_dao_tham_dinh") or "",
        ], size=9.5, center_cols=(0, 3, 4, 5))
    if not tasks_ky:
        tbl.add_row()
    _set_col_widths(tbl, [1, 5, 2.5, 2.5, 2.2, 2.5, 2.5])
    doc.add_paragraph()

    # ── 2. Nhiệm vụ tuần tới ──
    h2 = doc.add_paragraph()
    r = h2.add_run("2. Nhiệm vụ tuần tới ")
    r.bold = True; r.font.size = Pt(12)
    r2 = h2.add_run("(A/c lên dự kiến trước, có thể bổ sung thêm nhiệm vụ gửi e "
                     "trước 10h ngày thứ 2 tuần tới để e kịp tổng hợp, trình lãnh đạo "
                     "phòng đúng thời gian ạ)")
    r2.italic = True; r2.font.size = Pt(10.5)

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    _fill_row(tbl2.rows[0].cells, ["STT", "Nội dung", "Phối hợp (nếu có)", "Thời gian hoàn thành"], size=10)
    _style_header_row(tbl2.rows[0])
    for i, t in enumerate(nwt_rows, 1):
        cells = tbl2.add_row().cells
        _fill_row(cells, [str(i), t["noi_dung"], t.get("phoi_hop") or "",
                           t.get("thoi_gian_hoan_thanh") or ""],
                  size=9.5, center_cols=(0,))
    if not nwt_rows:
        tbl2.add_row()
    _set_col_widths(tbl2, [1, 8, 3, 4])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _danh_gia_tien_do(t):
    """So sánh Thời gian thực hiện hoàn thành với Thời hạn yêu cầu (thoi_han_vb).
    Nếu việc đã hoàn thành nhưng chưa ghi thời hạn, mặc định coi là đúng hạn
    (không thể kết luận là chậm khi không có mốc để so sánh)."""
    kt = t.get("ngay_kt") if t.get("trang_thai") == "Hoàn thành" else None
    if not kt:
        return ""
    han = t.get("thoi_han_vb")
    if not han:
        return "Đúng thời gian yêu cầu"
    try:
        d_han = datetime.date.fromisoformat(str(han))
        d_kt = datetime.date.fromisoformat(str(kt))
    except Exception:
        return "Đúng thời gian yêu cầu"
    if d_kt < d_han:
        return "Vượt thời gian yêu cầu"
    if d_kt == d_han:
        return "Đúng thời gian yêu cầu"
    return "Chậm so với yêu cầu"

def build_bieu01_report(ten, chuc_vu, ky_label, tasks, san_pham_intro, san_pham_list,
                         so_ngay_lam, so_ngay_nghi, mo_ta_them_gio, tu_danh_gia_pham_chat,
                         xep_loai, can_cu):
    """Tạo file .docx đúng theo bố cục Biểu 01: Cá nhân tự đánh giá."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)

    p0 = doc.add_paragraph(); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run("BIỂU 01: CÁ NHÂN TỰ ĐÁNH GIÁ\n")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(27, 58, 107)
    r = p0.add_run(f"Kết quả thực hiện công tác {ky_label}")
    r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph()

    p1 = doc.add_paragraph()
    r = p1.add_run("- Họ và tên: "); r.bold = True
    p1.add_run(ten).bold = True
    p2 = doc.add_paragraph()
    p2.add_run("- Chức vụ, đơn vị: ").bold = True
    p2.add_run(chuc_vu)
    doc.add_paragraph()

    tong = len(tasks)
    ht = sum(1 for t in tasks if t["trang_thai"] == "Hoàn thành")
    tx = sum(1 for t in tasks if (t.get("loai_viec_bc") or "Kế hoạch/thường xuyên") == "Kế hoạch/thường xuyên")
    ps = tong - tx
    danh_gia = [_danh_gia_tien_do(t) for t in tasks]
    dung_vuot = sum(1 for d in danh_gia if d in ("Đúng thời gian yêu cầu", "Vượt thời gian yêu cầu"))
    cham = sum(1 for d in danh_gia if d == "Chậm so với yêu cầu")
    pt = round(ht / tong * 100) if tong else 0

    h1 = doc.add_paragraph()
    h1.add_run(f"I- Kết quả thực hiện nhiệm vụ được giao trong {ky_label}").bold = True
    doc.add_paragraph(
        f"Trong {ky_label}, tham mưu, thực hiện {tong} công việc; trong đó: {tx} việc theo "
        f"chương trình, kế hoạch công tác và nhiệm vụ thường xuyên được giao; {ps} việc phát "
        f"sinh, đột xuất.")
    doc.add_paragraph(f"- Số việc hoàn thành: {ht}/{tong} việc, đạt tỷ lệ {pt}%.")
    doc.add_paragraph(
        f"- Tiến độ thực hiện: {dung_vuot} việc bảo đảm đúng thời gian yêu cầu; "
        f"chậm tiến độ: {cham if cham else 'Không'}.")

    if san_pham_list:
        pnb = doc.add_paragraph()
        pnb.add_run("Sản phẩm, giải pháp nổi bật, có tính đổi mới sáng tạo trong kỳ:").bold = True
        pnb.runs[0].italic = True
        if san_pham_intro:
            doc.add_paragraph(san_pham_intro)
        for i, sp in enumerate(san_pham_list, 1):
            doc.add_paragraph(f"({i}) {sp}")
    doc.add_paragraph()

    # ── Bảng chi tiết (trang ngang) ──
    sec_landscape = doc.add_section()
    sec_landscape.orientation = 1  # WD_ORIENT.LANDSCAPE
    sec_landscape.page_width, sec_landscape.page_height = sec_landscape.page_height, sec_landscape.page_width
    sec_landscape.left_margin = Cm(1.5); sec_landscape.right_margin = Cm(1.5)
    sec_landscape.top_margin = Cm(1.5); sec_landscape.bottom_margin = Cm(1.5)

    tbl = doc.add_table(rows=1, cols=9)
    tbl.style = "Table Grid"
    heads = ["STT", "Nội dung công việc được giao/ thực hiện", "Chủ trì", "Phối hợp",
              "Thời hạn yêu cầu hoàn thành", "Thời gian thực hiện hoàn thành",
              "Đánh giá tiến độ thực hiện", "Tự đánh giá chất lượng công tác tham mưu",
              "Loại việc"]
    _fill_row(tbl.rows[0].cells, heads, size=9.5)
    _style_header_row(tbl.rows[0])
    for i, t in enumerate(tasks, 1):
        cells = tbl.add_row().cells
        kt = _fmt_d(t.get("ngay_kt")) if t.get("trang_thai") == "Hoàn thành" else ""
        _fill_row(cells, [
            str(i), t["title"],
            "X" if t.get("chu_tri", True) else "",
            t.get("phoi_hop") or "",
            _fmt_d(t.get("thoi_han_vb")),
            kt,
            _danh_gia_tien_do(t),
            t.get("tu_danh_gia_cl") or ("Hoàn thành tốt nhiệm vụ" if t["trang_thai"] == "Hoàn thành" else ""),
            t.get("loai_viec_bc") or "Kế hoạch/thường xuyên",
        ], size=9, center_cols=(0, 2, 4, 5, 6, 8))
    if not tasks:
        tbl.add_row()
    _set_col_widths(tbl, [1, 6, 1.5, 2.5, 2.3, 2.3, 2.5, 4, 2.5])

    # ── Quay lại trang dọc cho mục II, III ──
    sec_portrait = doc.add_section()
    sec_portrait.orientation = 0  # WD_ORIENT.PORTRAIT
    sec_portrait.page_width, sec_portrait.page_height = sec_portrait.page_height, sec_portrait.page_width
    sec_portrait.left_margin = Cm(2.5); sec_portrait.right_margin = Cm(2)
    sec_portrait.top_margin = Cm(2); sec_portrait.bottom_margin = Cm(2)

    h2 = doc.add_paragraph()
    h2.add_run("II- Đánh giá về đạo đức, lối sống; việc chấp hành nội quy, quy chế làm việc").bold = True
    doc.add_paragraph(f"- Số ngày làm việc: {so_ngay_lam}; số ngày nghỉ: {so_ngay_nghi}. {mo_ta_them_gio}")
    doc.add_paragraph(f"- Tự đánh giá ngắn gọn về phẩm chất, đạo đức, lối sống: {tu_danh_gia_pham_chat}")
    doc.add_paragraph()

    h3 = doc.add_paragraph()
    h3.add_run(f"III. Đề xuất mức xếp loại (theo 4 mức): {xep_loai}.").bold = True
    p_cc = doc.add_paragraph()
    r = p_cc.add_run(f"Căn cứ đề xuất: {can_cu}")
    r.italic = True
    doc.add_paragraph()

    kytbl = doc.add_table(rows=1, cols=2)
    kytbl.autofit = True
    kytbl.rows[0].cells[0].text = ""
    c2 = kytbl.rows[0].cells[1]
    c2.text = ""
    pk = c2.paragraphs[0]
    pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pk.add_run("NGƯỜI TỰ ĐÁNH GIÁ\n"); r.bold = True
    pk.add_run("\n\n\n")
    r2 = pk.add_run(ten); r2.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════
#  BÁO CÁO TỔNG HỢP THEO KỲ (kế thừa bản desktop)
# ══════════════════════════════════════════════════════
XL_MUC = ["Tự động (theo kết quả)", "Hoàn thành xuất sắc nhiệm vụ",
          "Hoàn thành tốt nhiệm vụ", "Hoàn thành nhiệm vụ", "Không hoàn thành nhiệm vụ"]

def ky_options():
    today = datetime.date.today()
    q = (today.month - 1) // 3 + 1
    q_roman = {1: "I", 2: "II", 3: "III", 4: "IV"}[q]
    return {
        "Tuần": (today - datetime.timedelta(days=today.weekday()),
                 today - datetime.timedelta(days=today.weekday()) + datetime.timedelta(days=6),
                 f"Tuần {today.isocalendar()[1]}/{today.year}"),
        "Tháng": (today.replace(day=1),
                  today.replace(day=calendar.monthrange(today.year, today.month)[1]),
                  f"Tháng {today.month:02d}/{today.year}"),
        "Quý": (datetime.date(today.year, (q - 1) * 3 + 1, 1),
                datetime.date(today.year, q * 3, calendar.monthrange(today.year, q * 3)[1]),
                f"Quý {q_roman}/{today.year}"),
        "6 tháng đầu năm": (datetime.date(today.year, 1, 1), datetime.date(today.year, 6, 30), f"6 tháng đầu năm {today.year}"),
        "9 tháng đầu năm": (datetime.date(today.year, 1, 1), datetime.date(today.year, 9, 30), f"9 tháng đầu năm {today.year}"),
        "Năm": (datetime.date(today.year, 1, 1), datetime.date(today.year, 12, 31), f"Năm {today.year}"),
        "Tùy chọn": None,
    }

def chon_ky(prefix, today):
    """Hiển thị bộ chọn kỳ báo cáo dùng chung; trả về (tu, den, nhãn kỳ)."""
    options = ky_options()
    ky_chon = st.selectbox("Chọn kỳ báo cáo", list(options.keys()), key=f"{prefix}_ky")
    if ky_chon == "Tùy chọn":
        c1, c2, c3 = st.columns(3)
        tu = c1.date_input("Từ ngày", value=today.replace(day=1), key=f"{prefix}_tu")
        den = c2.date_input("Đến ngày", value=today, key=f"{prefix}_den")
        td = c3.text_input("Tiêu đề / nhãn kỳ", value="kỳ báo cáo", key=f"{prefix}_td")
    else:
        tu, den, td = options[ky_chon]
    return tu, den, td

def tinh_toan_ky(tu, den, xl_chon):
    tasks = db_query(tu, den)
    tong = len(tasks)
    ht = sum(1 for t in tasks if t["trang_thai"] == "Hoàn thành")
    tre = sum(1 for t in tasks if t["trang_thai"] == "Trễ hạn")
    dang = sum(1 for t in tasks if t["trang_thai"] == "Đang thực hiện")
    pt = round(ht / tong * 100) if tong else 0
    theo_loai = Counter(t.get("loai") or "Khác" for t in tasks)
    noi_bat = [t for t in tasks if t.get("loai") == "Xây dựng ứng dụng chuyển đổi số"
               and t["trang_thai"] == "Hoàn thành"]
    if xl_chon == XL_MUC[0]:
        if tong == 0:
            xl = "Chưa có dữ liệu để đánh giá"
        elif pt == 100 and tre == 0 and noi_bat:
            xl = "Hoàn thành xuất sắc nhiệm vụ"
        elif pt >= 90 and tre == 0:
            xl = "Hoàn thành tốt nhiệm vụ"
        elif pt >= 70:
            xl = "Hoàn thành nhiệm vụ"
        else:
            xl = "Không hoàn thành nhiệm vụ"
    else:
        xl = xl_chon
    can_cu = (f"hoàn thành {ht}/{tong} công việc được giao, đạt tỷ lệ {pt}%"
              + (f"; còn {tre} việc trễ hạn" if tre else "; không có việc trễ hạn")
              + (f". Trong kỳ có {len(noi_bat)} sản phẩm/giải pháp ứng dụng công nghệ "
                 "thông tin được xây dựng và đưa vào sử dụng hiệu quả." if noi_bat else "."))
    return dict(tasks=tasks, tong=tong, ht=ht, tre=tre, dang=dang, pt=pt,
                theo_loai=theo_loai, noi_bat=noi_bat, xep_loai=xl, can_cu=can_cu)

def build_ky_report(ten, chuc_vu, tu, den, td, d, nhan_xet, phuong_huong):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)

    p0 = doc.add_paragraph(); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run("ĐẢNG CỘNG SẢN VIỆT NAM\n"); r.bold = True; r.font.size = Pt(13)
    r = p0.add_run("TỈNH ỦY TUYÊN QUANG\n"); r.bold = True; r.font.size = Pt(12)
    r = p0.add_run("Ban Tuyên giáo và Dân vận\n"); r.font.size = Pt(11)
    doc.add_paragraph()

    pt2 = doc.add_paragraph(); pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pt2.add_run(f"BÁO CÁO CÔNG TÁC {td.upper()}")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(27, 58, 107)

    p_sub = doc.add_paragraph(); p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run(f"(Từ ngày {_fmt_d(tu)} đến ngày {_fmt_d(den)})").font.size = Pt(11)
    doc.add_paragraph()

    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.add_run(f"Người lập: {ten}\nChức vụ: {chuc_vu}\n"
               f"Ngày lập: {datetime.date.today().strftime('%d/%m/%Y')}").font.size = Pt(10)
    doc.add_paragraph()

    h = doc.add_heading("I. TỔNG KẾT KẾT QUẢ", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    _fill_row(tbl.rows[0].cells, ["Chỉ tiêu", "Số lượng", "Tỷ lệ"], size=11)
    _style_header_row(tbl.rows[0])
    for row_data in [("Tổng số công việc", str(d["tong"]), "100%"),
                      ("Đã hoàn thành", str(d["ht"]), f"{d['pt']}%"),
                      ("Đang thực hiện", str(d["dang"]), ""),
                      ("Trễ hạn", str(d["tre"]), "")]:
        cells = tbl.add_row().cells
        _fill_row(cells, row_data, center_cols=(1, 2))
    _set_col_widths(tbl, [9, 4, 4])
    doc.add_paragraph()

    if d["theo_loai"]:
        doc.add_paragraph("Phân loại theo tính chất công việc:").runs[0].bold = True
        tbl2 = doc.add_table(rows=1, cols=2); tbl2.style = "Table Grid"
        _fill_row(tbl2.rows[0].cells, ["Loại công việc", "Số lượng"], size=11)
        _style_header_row(tbl2.rows[0])
        for k, v in sorted(d["theo_loai"].items(), key=lambda x: -x[1]):
            cells = tbl2.add_row().cells
            _fill_row(cells, [k, str(v)], center_cols=(1,))
        _set_col_widths(tbl2, [13, 4])
        doc.add_paragraph()

    h = doc.add_heading("II. CHI TIẾT CÔNG VIỆC", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
    tbl3 = doc.add_table(rows=1, cols=7); tbl3.style = "Table Grid"
    heads = ["STT", "Ngày", "Nội dung công việc", "Loại việc", "Thời gian", "Trạng thái", "Kết quả"]
    _fill_row(tbl3.rows[0].cells, heads, size=10.5)
    _style_header_row(tbl3.rows[0])
    for i, t in enumerate(d["tasks"], 1):
        gio = f"{t.get('gio_bd','') or ''}-{t.get('gio_kt','') or ''}".strip("-") or "Cả ngày"
        cells = tbl3.add_row().cells
        _fill_row(cells, [str(i), _fmt_d(t["ngay_bd"]), t["title"], t.get("loai", ""),
                           gio, t["trang_thai"], t.get("ket_qua", "") or ""],
                  size=10, center_cols=(0, 1, 4, 5))
    _set_col_widths(tbl3, [1, 2.2, 6, 3, 2.2, 2.3, 2.3])
    doc.add_paragraph()

    if d["noi_bat"]:
        h = doc.add_heading("III. SẢN PHẨM, GIẢI PHÁP NỔI BẬT TRONG KỲ", level=2)
        for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
        for t in d["noi_bat"]:
            pb = doc.add_paragraph(style="List Bullet")
            r = pb.add_run(t["title"]); r.bold = True
            if t.get("ket_qua"):
                pb.add_run(f" — {t['ket_qua']}")
        doc.add_paragraph()

    h = doc.add_heading("IV. NHẬN XÉT, ĐÁNH GIÁ", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
    doc.add_paragraph(f"- Hoàn thành {d['ht']}/{d['tong']} công việc, đạt tỷ lệ {d['pt']}%.")
    if d["tre"]:
        doc.add_paragraph(f"- Còn {d['tre']} việc trễ hạn, cần rút kinh nghiệm về tiến độ.")
    if nhan_xet:
        doc.add_paragraph(f"- {nhan_xet}")
    doc.add_paragraph()

    h = doc.add_heading("V. PHƯƠNG HƯỚNG NHIỆM VỤ TIẾP THEO", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
    if phuong_huong:
        for line in phuong_huong.split("\n"):
            if line.strip():
                doc.add_paragraph(f"- {line.strip()}")
    else:
        doc.add_paragraph("- Tiếp tục thực hiện các nhiệm vụ được phân công.")
    doc.add_paragraph()

    h = doc.add_heading("VI. ĐỀ XUẤT MỨC XẾP LOẠI", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(27, 58, 107)
    p_xl = doc.add_paragraph(); r = p_xl.add_run(d["xep_loai"]); r.bold = True; r.font.size = Pt(12)
    p_cc = doc.add_paragraph(); r2 = p_cc.add_run(f"Căn cứ đề xuất: {d['can_cu']}")
    r2.italic = True; r2.font.size = Pt(10.5)
    doc.add_paragraph()

    ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today = datetime.date.today()
    ps.add_run(f"Tuyên Quang, ngày {today.day} tháng {today.month} năm {today.year}\n").font.size = Pt(11)
    r2 = ps.add_run("NGƯỜI LẬP BÁO CÁO\n\n\n\n"); r2.bold = True
    r3 = ps.add_run(ten); r3.bold = True

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════
ten = cfg_get("ten", "Chưa đặt tên")
chuc_vu = cfg_get("chuc_vu", "Chuyên viên")
chuc_vu_day_du = cfg_get("chuc_vu_day_du", "") or chuc_vu

with st.sidebar:
    st.markdown(f"### ⭐ TỈNH ỦY TUYÊN QUANG")
    st.caption("Ban Tuyên giáo & Dân vận")
    st.markdown(f"**👤 {ten}**")
    st.caption(chuc_vu)
    st.divider()
    page = st.radio("Điều hướng", [
        "📊 Dashboard", "📅 Lịch & Quản lý công việc",
        "🗓️ Nhiệm vụ tuần tới", "📄 Báo cáo & Xuất file", "⚙️ Cài đặt"
    ], label_visibility="collapsed")
    st.divider()
    st.caption("v3.0 (Web) © 2026 TGDV Tuyên Quang")

# ══════════════════════════════════════════════════════
#  TRANG: DASHBOARD
# ══════════════════════════════════════════════════════
if page == "📊 Dashboard":
    st.title(f"📊 Xin chào, {ten}!")
    st.caption(datetime.date.today().strftime("%A, ngày %d tháng %m năm %Y"))

    tasks = db_all()
    tong = len(tasks)
    ht = sum(1 for t in tasks if t["trang_thai"] == "Hoàn thành")
    dang = sum(1 for t in tasks if t["trang_thai"] == "Đang thực hiện")
    chua = sum(1 for t in tasks if t["trang_thai"] == "Chưa thực hiện")
    tre = sum(1 for t in tasks if t["trang_thai"] == "Trễ hạn")
    pt = round(ht / tong * 100) if tong else 0

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("📋 Tổng công việc", tong)
    c2.metric("✅ Hoàn thành", f"{ht} ({pt}%)")
    c3.metric("🔄 Đang thực hiện", dang)
    c4.metric("⏳ Chưa thực hiện", chua)
    c5.metric("⚠️ Trễ hạn", tre)

    col1, col2 = st.columns([1, 2])
    with col1:
        st.subheader("🎯 Trạng thái công việc")
        if tasks:
            df = pd.DataFrame(Counter(t["trang_thai"] for t in tasks).items(), columns=["Trạng thái", "Số lượng"])
            fig = px.pie(df, names="Trạng thái", values="Số lượng", hole=0.4,
                         color="Trạng thái",
                         color_discrete_map={"Hoàn thành": GREEN, "Đang thực hiện": ORANGE,
                                              "Chưa thực hiện": "#7D3C98", "Trễ hạn": DANGER})
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Chưa có dữ liệu")
    with col2:
        st.subheader(f"📊 Tiến độ theo tháng — {datetime.date.today().year}")
        y = datetime.date.today().year
        rows = []
        for m in range(1, 13):
            fd = f"{y}-{m:02d}-01"
            ld = f"{y}-{m:02d}-{calendar.monthrange(y, m)[1]:02d}"
            mt = db_query(fd, ld)
            rows.append({"Tháng": f"T{m}", "Tổng CV": len(mt),
                         "Hoàn thành": sum(1 for t in mt if t["trang_thai"] == "Hoàn thành")})
        df = pd.DataFrame(rows).melt(id_vars="Tháng", var_name="Loại", value_name="Số lượng")
        fig = px.bar(df, x="Tháng", y="Số lượng", color="Loại", barmode="group",
                     color_discrete_map={"Tổng CV": "#2A5298", "Hoàn thành": GREEN})
        st.plotly_chart(fig, use_container_width=True)

    col3, col4 = st.columns(2)
    with col3:
        st.subheader("🏷️ Phân loại công việc")
        if tasks:
            df = pd.DataFrame(Counter(t.get("loai") or "Khác" for t in tasks).items(),
                               columns=["Loại", "Số lượng"]).sort_values("Số lượng")
            fig = px.bar(df, x="Số lượng", y="Loại", orientation="h", color="Loại")
            fig.update_layout(showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
    with col4:
        st.subheader("📈 Xu hướng 8 tuần gần nhất")
        today = datetime.date.today()
        rows = []
        for w in range(7, -1, -1):
            mon = today - datetime.timedelta(days=today.weekday()) - datetime.timedelta(weeks=w)
            sun = mon + datetime.timedelta(days=6)
            wt = db_query(str(mon), str(sun))
            rows.append({"Tuần": f"T{mon.isocalendar()[1]}", "Tổng CV": len(wt),
                         "Hoàn thành": sum(1 for t in wt if t["trang_thai"] == "Hoàn thành")})
        df = pd.DataFrame(rows).melt(id_vars="Tuần", var_name="Loại", value_name="Số lượng")
        fig = px.line(df, x="Tuần", y="Số lượng", color="Loại", markers=True,
                      color_discrete_map={"Tổng CV": "#2A5298", "Hoàn thành": GREEN})
        st.plotly_chart(fig, use_container_width=True)

    st.subheader(f"📌 Công việc hôm nay")
    ds = str(datetime.date.today())
    today_tasks = db_query(ds, ds)
    if not today_tasks:
        st.info("✨ Không có công việc nào hôm nay")
    else:
        for t in today_tasks:
            gio = f"{t.get('gio_bd','') or ''} – {t.get('gio_kt','') or ''}".strip(" –") or "Cả ngày"
            st.markdown(f"**[{gio}]** {t['title']} — *{t['trang_thai']}*")

# ══════════════════════════════════════════════════════
#  TRANG: LỊCH & QUẢN LÝ CÔNG VIỆC
# ══════════════════════════════════════════════════════
elif page == "📅 Lịch & Quản lý công việc":
    st.title("📅 Lịch & Quản lý công việc")

    with st.expander("➕ Thêm / ✏️ Sửa công việc", expanded=bool(st.session_state.get("edit_task_id"))):
        edit_id = st.session_state.get("edit_task_id")
        edit_data = {}
        if edit_id:
            r = sb.table("tasks").select("*").eq("id", edit_id).execute().data
            edit_data = r[0] if r else {}
            st.info(f"Đang sửa: {edit_data.get('title','')}")

        c1, c2 = st.columns(2)
        title = c1.text_input("Tiêu đề (*)", value=edit_data.get("title", ""))
        loai = c2.selectbox("Loại công việc", LOAI_CV,
                             index=LOAI_CV.index(edit_data["loai"]) if edit_data.get("loai") in LOAI_CV else 0)

        c3, c4, c5 = st.columns(3)
        ngay_bd = c3.date_input("Từ ngày", value=datetime.date.fromisoformat(edit_data["ngay_bd"]) if edit_data.get("ngay_bd") else datetime.date.today())
        ngay_kt = c4.date_input("Đến ngày", value=datetime.date.fromisoformat(edit_data["ngay_kt"]) if edit_data.get("ngay_kt") else ngay_bd)
        uu_tien = c5.selectbox("Ưu tiên", UU_TIEN, index=UU_TIEN.index(edit_data["uu_tien"]) if edit_data.get("uu_tien") in UU_TIEN else 1)

        c6, c7, c8 = st.columns(3)
        gio_bd = c6.text_input("Giờ bắt đầu", value=edit_data.get("gio_bd") or "07:30")
        gio_kt = c7.text_input("Giờ kết thúc", value=edit_data.get("gio_kt") or "17:00")
        trang_thai = c8.selectbox("Trạng thái", TRANG_THAI,
                                   index=TRANG_THAI.index(edit_data["trang_thai"]) if edit_data.get("trang_thai") in TRANG_THAI else 0)

        c9, c10 = st.columns(2)
        dia_diem = c9.text_input("Địa điểm", value=edit_data.get("dia_diem") or "")
        nguoi_pt = c10.text_input("Người phụ trách", value=edit_data.get("nguoi_pt") or "")

        st.markdown("**Thông tin cho Mẫu báo cáo CBCC VP** (không bắt buộc)")
        c11, c12, c13 = st.columns(3)
        lanh_dao_giao = c11.text_input("Lãnh đạo giao", value=edit_data.get("lanh_dao_giao") or "")
        thoi_han_vb_raw = edit_data.get("thoi_han_vb")
        thoi_han_vb = c12.date_input("Thời hạn yêu cầu hoàn thành / Thời hạn của VB",
                                      value=datetime.date.fromisoformat(thoi_han_vb_raw) if thoi_han_vb_raw else None)
        lanh_dao_tham_dinh = c13.text_input("Lãnh đạo thẩm định, duyệt", value=edit_data.get("lanh_dao_tham_dinh") or "")

        st.markdown("**Thông tin cho Biểu 01 — Phiếu tự đánh giá** (không bắt buộc)")
        c14, c15, c16 = st.columns(3)
        chu_tri = c14.checkbox("Chủ trì thực hiện", value=edit_data.get("chu_tri", True) if edit_data else True)
        phoi_hop = c15.text_input("Phối hợp (đơn vị/người, nếu có)", value=edit_data.get("phoi_hop") or "")
        loai_viec_bc = c16.selectbox("Loại việc (cho Biểu 01)", LOAI_VIEC_BC,
                                      index=LOAI_VIEC_BC.index(edit_data["loai_viec_bc"]) if edit_data.get("loai_viec_bc") in LOAI_VIEC_BC else 0)
        tu_danh_gia_cl = st.text_input("Tự đánh giá chất lượng công tác tham mưu",
                                        value=edit_data.get("tu_danh_gia_cl") or "",
                                        placeholder="VD: Hoàn thành tốt nhiệm vụ")

        ket_qua = st.text_input("Kết quả thực hiện", value=edit_data.get("ket_qua") or "")
        ghi_chu = st.text_area("Ghi chú", value=edit_data.get("ghi_chu") or "", height=70)

        bc1, bc2, bc3 = st.columns([1, 1, 4])
        if bc1.button("💾 Lưu", type="primary"):
            if not title.strip():
                st.warning("Vui lòng nhập tiêu đề!")
            else:
                data = dict(title=title.strip(), loai=loai, uu_tien=uu_tien,
                            ngay_bd=str(ngay_bd), ngay_kt=str(ngay_kt),
                            gio_bd=gio_bd or None, gio_kt=gio_kt or None,
                            dia_diem=dia_diem or None, nguoi_pt=nguoi_pt or None,
                            trang_thai=trang_thai, ket_qua=ket_qua or None,
                            ghi_chu=ghi_chu or None,
                            lanh_dao_giao=lanh_dao_giao or None,
                            thoi_han_vb=str(thoi_han_vb) if thoi_han_vb else None,
                            lanh_dao_tham_dinh=lanh_dao_tham_dinh or None,
                            chu_tri=chu_tri, phoi_hop=phoi_hop or None,
                            loai_viec_bc=loai_viec_bc,
                            tu_danh_gia_cl=tu_danh_gia_cl or None)
                if edit_id:
                    db_sua(edit_id, data)
                    st.session_state["edit_task_id"] = None
                else:
                    db_them(data)
                st.success("Đã lưu!")
                st.rerun()
        if edit_id and bc2.button("✖ Hủy sửa"):
            st.session_state["edit_task_id"] = None
            st.rerun()

    st.divider()
    fc1, fc2, fc3, fc4 = st.columns([2, 1.5, 1.5, 1])
    kw = fc1.text_input("🔍 Tìm kiếm")
    f_ts = fc2.selectbox("Trạng thái", ["Tất cả"] + TRANG_THAI)
    f_loai = fc3.selectbox("Loại", ["Tất cả"] + LOAI_CV)
    date_range = fc4.date_input("Khoảng ngày", value=(), help="Để trống = tất cả")

    if isinstance(date_range, tuple) and len(date_range) == 2:
        tasks = db_query(str(date_range[0]), str(date_range[1]))
    else:
        tasks = db_all()
    if kw:
        tasks = [t for t in tasks if kw.lower() in t["title"].lower() or kw.lower() in (t.get("dia_diem") or "").lower()]
    if f_ts != "Tất cả":
        tasks = [t for t in tasks if t["trang_thai"] == f_ts]
    if f_loai != "Tất cả":
        tasks = [t for t in tasks if t.get("loai") == f_loai]

    st.caption(f"📋 Tổng: {len(tasks)} công việc")
    for t in tasks:
        gio = f"{t.get('gio_bd','') or ''} – {t.get('gio_kt','') or ''}".strip(" –") or "Cả ngày"
        icon = {"Hoàn thành": "✅", "Đang thực hiện": "🔄", "Trễ hạn": "⚠️"}.get(t["trang_thai"], "⏳")
        with st.container(border=True):
            cc1, cc2, cc3 = st.columns([5, 2, 1.5])
            cc1.markdown(f"{icon} **{t['title']}**  \n"
                         f"📅 {_fmt_d(t['ngay_bd'])} → {_fmt_d(t['ngay_kt'])}  ·  ⏰ {gio}  ·  🏷️ {t.get('loai','')}")
            cc2.markdown(f"**{t['trang_thai']}**  \n📍 {t.get('dia_diem') or '—'}")
            b1, b2, b3 = cc3.columns(3)
            if b1.button("✏️", key=f"e{t['id']}"):
                st.session_state["edit_task_id"] = t["id"]
                st.rerun()
            if b2.button("✅", key=f"h{t['id']}"):
                db_sua(t["id"], {"trang_thai": "Hoàn thành"})
                st.rerun()
            if b3.button("🗑", key=f"x{t['id']}"):
                db_xoa(t["id"])
                st.rerun()

# ══════════════════════════════════════════════════════
#  TRANG: NHIỆM VỤ TUẦN TỚI
# ══════════════════════════════════════════════════════
elif page == "🗓️ Nhiệm vụ tuần tới":
    st.title("🗓️ Nhiệm vụ tuần tới")
    st.caption("Dùng cho mục 2 của Mẫu Tổng hợp BC của CBCC VP — lên kế hoạch trước cho tuần sau.")

    today = datetime.date.today()
    next_monday = today + datetime.timedelta(days=(7 - today.weekday()))
    tuan_mac_dinh = next_monday.isocalendar()[1]
    nam_mac_dinh = next_monday.isocalendar()[0]

    c1, c2 = st.columns(2)
    tuan_chon = c1.number_input("Tuần (số ISO)", min_value=1, max_value=53, value=tuan_mac_dinh)
    nam_chon = c2.number_input("Năm", min_value=2020, max_value=2100, value=nam_mac_dinh)

    with st.expander("➕ Thêm nhiệm vụ dự kiến", expanded=True):
        noi_dung = st.text_input("Nội dung (*)")
        cc1, cc2 = st.columns(2)
        phoi_hop = cc1.text_input("Phối hợp (nếu có)")
        tg_hoan_thanh = cc2.text_input("Thời gian hoàn thành dự kiến", placeholder="VD: Trước 15/07")
        if st.button("💾 Thêm vào tuần tới", type="primary"):
            if not noi_dung.strip():
                st.warning("Vui lòng nhập nội dung!")
            else:
                nwt_them({"noi_dung": noi_dung.strip(), "phoi_hop": phoi_hop or None,
                           "thoi_gian_hoan_thanh": tg_hoan_thanh or None,
                           "tuan": int(tuan_chon), "nam": int(nam_chon)})
                st.success("Đã thêm!")
                st.rerun()

    st.divider()
    rows = nwt_query(int(tuan_chon), int(nam_chon))
    st.caption(f"📋 Tuần {tuan_chon}/{nam_chon}: {len(rows)} nhiệm vụ dự kiến")
    for i, r in enumerate(rows, 1):
        with st.container(border=True):
            cc1, cc2 = st.columns([5, 1])
            cc1.markdown(f"**{i}. {r['noi_dung']}**  \n"
                         f"🤝 Phối hợp: {r.get('phoi_hop') or '—'}  ·  ⏰ {r.get('thoi_gian_hoan_thanh') or '—'}")
            if cc2.button("🗑 Xóa", key=f"nwt{r['id']}"):
                nwt_xoa(r["id"])
                st.rerun()

# ══════════════════════════════════════════════════════
#  TRANG: BÁO CÁO & XUẤT FILE
# ══════════════════════════════════════════════════════
elif page == "📄 Báo cáo & Xuất file":
    st.title("📄 Báo cáo & Xuất file")

    tab1, tab2, tab3 = st.tabs([
        "📝 Mẫu Tổng hợp BC của CBCC VP (theo tuần)",
        "🗂️ Biểu 01: Phiếu tự đánh giá",
        "📊 Báo cáo tổng hợp theo kỳ",
    ])

    # ── TAB 1: Mẫu CBCC VP ──
    with tab1:
        st.caption("Xuất đúng bố cục Mẫu - Tổng hợp BC của CBCC VP: mục 1 lấy công việc thực hiện "
                   "trong tuần đã chọn, mục 2 lấy danh sách ở trang **Nhiệm vụ tuần tới**.")
        today = datetime.date.today()
        c1, c2 = st.columns(2)
        mon_report = c1.date_input("Tuần báo cáo — chọn 1 ngày trong tuần", value=today)
        mon_start = mon_report - datetime.timedelta(days=mon_report.weekday())
        mon_end = mon_start + datetime.timedelta(days=6)
        c2.markdown(f"**Khoảng ngày:** {mon_start.strftime('%d/%m/%Y')} – {mon_end.strftime('%d/%m/%Y')}  \n"
                    f"**Tuần ISO:** {mon_start.isocalendar()[1]}/{mon_start.isocalendar()[0]}")

        tasks_ky = db_query(str(mon_start), str(mon_end))
        next_mon = mon_start + datetime.timedelta(days=7)
        nwt_rows = nwt_query(next_mon.isocalendar()[1], next_mon.isocalendar()[0])
        st.info(f"Mục 1 sẽ có **{len(tasks_ky)}** công việc · Mục 2 sẽ có **{len(nwt_rows)}** "
                f"nhiệm vụ (tuần {next_mon.isocalendar()[1]}/{next_mon.isocalendar()[0]})")

        if st.button("⬇ Xuất file Word theo Mẫu CBCC VP", type="primary"):
            buf = build_cbcc_report(ten, tasks_ky, nwt_rows, mon_start.isocalendar()[1], mon_start.isocalendar()[0])
            st.download_button("📥 Tải file .docx", data=buf,
                                file_name=f"BC_CBCC_Tuan{mon_start.isocalendar()[1]}_{mon_start.isocalendar()[0]}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # ── TAB 2: Biểu 01 — Phiếu tự đánh giá ──
    with tab2:
        st.caption("Xuất đúng bố cục **Biểu 01: Cá nhân tự đánh giá** — dùng được cho báo cáo "
                   "tuần, tháng, quý, 6 tháng, 9 tháng hoặc năm. Mục I và bảng chi tiết được "
                   "tổng hợp tự động từ công việc trong kỳ; mục II (đạo đức, lối sống) bạn tự "
                   "viết vì đây là nội dung mang tính tường thuật cá nhân.")
        today = datetime.date.today()
        tu, den, td = chon_ky("bieu01", today)
        st.markdown(f"**Khoảng ngày:** {_fmt_d(tu)} – {_fmt_d(den)}")

        tasks_ky = db_query(tu, den)
        ht_ky = [t for t in tasks_ky if t["trang_thai"] == "Hoàn thành"]
        noi_bat = [t for t in tasks_ky if t.get("loai") == "Xây dựng ứng dụng chuyển đổi số"
                   and t["trang_thai"] == "Hoàn thành"]
        st.info(f"Mục I / bảng chi tiết sẽ có **{len(tasks_ky)}** công việc, trong đó "
                f"**{len(noi_bat)}** việc được nhận diện là sản phẩm/giải pháp nổi bật "
                f"(loại = *Xây dựng ứng dụng chuyển đổi số* và đã hoàn thành).")

        if tasks_ky:
            st.markdown("**Chỉnh nhanh các cột riêng cho Biểu 01** (Thời hạn yêu cầu, Chủ trì, "
                       "Phối hợp, Tự đánh giá chất lượng, Loại việc) — sửa xong bấm Lưu, "
                       "không cần vào từng công việc ở trang Quản lý công việc.")
            edit_df = pd.DataFrame([{
                "id": t["id"],
                "Nội dung": t["title"],
                "Trạng thái": t["trang_thai"],
                "Chủ trì": t.get("chu_tri", True) if t.get("chu_tri") is not None else True,
                "Phối hợp": t.get("phoi_hop") or "",
                "Thời hạn yêu cầu": datetime.date.fromisoformat(t["thoi_han_vb"]) if t.get("thoi_han_vb") else None,
                "Tự đánh giá chất lượng": t.get("tu_danh_gia_cl") or "",
                "Loại việc": t.get("loai_viec_bc") or "Kế hoạch/thường xuyên",
            } for t in tasks_ky])
            edited = st.data_editor(
                edit_df, hide_index=True, use_container_width=True, key="bieu01_editor",
                disabled=["id", "Nội dung", "Trạng thái"],
                column_config={
                    "id": None,  # ẩn cột id khỏi bảng hiển thị
                    "Loại việc": st.column_config.SelectboxColumn(options=LOAI_VIEC_BC),
                    "Thời hạn yêu cầu": st.column_config.DateColumn(format="DD/MM/YYYY"),
                })
            if st.button("💾 Lưu các chỉnh sửa vào công việc"):
                for _, row in edited.iterrows():
                    db_sua(int(row["id"]), {
                        "chu_tri": bool(row["Chủ trì"]),
                        "phoi_hop": row["Phối hợp"] or None,
                        "thoi_han_vb": str(row["Thời hạn yêu cầu"]) if pd.notna(row["Thời hạn yêu cầu"]) else None,
                        "tu_danh_gia_cl": row["Tự đánh giá chất lượng"] or None,
                        "loai_viec_bc": row["Loại việc"],
                    })
                st.success("Đã lưu! Số liệu và bảng chi tiết bên dưới đã được cập nhật.")
                st.rerun()

        st.markdown("**Sản phẩm, giải pháp nổi bật** (mục I)")
        san_pham_intro = st.text_area("Câu dẫn (không bắt buộc)",
            value=("Ngoài khối lượng công việc thường xuyên, trong kỳ đã chủ động nghiên cứu, "
                   "tự xây dựng và đưa vào sử dụng thực tế các sản phẩm, phần mềm phục vụ trực "
                   "tiếp công tác chuyên môn và chuyển đổi số của Ban:") if noi_bat else "",
            height=70)
        san_pham_list = st.text_area(
            "Danh sách sản phẩm — mỗi dòng 1 sản phẩm (có thể sửa/thêm)",
            value="\n".join(t["title"] + (f" — {t['ket_qua']}" if t.get("ket_qua") else "") for t in noi_bat),
            height=100)
        san_pham_lines = [l.strip() for l in san_pham_list.split("\n") if l.strip()]

        st.markdown("**II- Đánh giá về đạo đức, lối sống; chấp hành nội quy, quy chế làm việc**")
        c1, c2 = st.columns(2)
        so_ngay_lam = c1.text_input("Số ngày làm việc", placeholder="VD: khoảng 63 ngày công")
        so_ngay_nghi = c2.text_input("Số ngày nghỉ", placeholder="VD: Không")
        mo_ta_them_gio = st.text_area("Mô tả thêm về thời gian làm việc (làm ngoài giờ, cuối tuần...)", height=70)
        tu_danh_gia_pham_chat = st.text_area("Tự đánh giá phẩm chất, đạo đức, lối sống", height=90)

        xl_chon2 = st.selectbox("Đề xuất mức xếp loại", XL_MUC, key="bieu01_xl")
        d2 = tinh_toan_ky(tu, den, xl_chon2)
        c3, c4, c5 = st.columns(3)
        c3.metric("Tổng công việc", d2["tong"])
        c4.metric("Hoàn thành", f"{d2['ht']} ({d2['pt']}%)")
        c5.metric("Đề xuất xếp loại", d2["xep_loai"])

        if st.button("⬇ Xuất file Word theo Biểu 01", type="primary"):
            if not so_ngay_lam or not tu_danh_gia_pham_chat:
                st.warning("Vui lòng điền ít nhất Số ngày làm việc và Tự đánh giá phẩm chất, đạo đức trước khi xuất.")
            else:
                buf = build_bieu01_report(ten, chuc_vu_day_du, td, tasks_ky, san_pham_intro, san_pham_lines,
                                           so_ngay_lam, so_ngay_nghi or "Không", mo_ta_them_gio,
                                           tu_danh_gia_pham_chat, d2["xep_loai"], d2["can_cu"])
                st.download_button("📥 Tải file .docx", data=buf,
                                    file_name=f"Bieu01_TuDanhGia_{td.replace(' ', '_').replace('/', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # ── TAB 3: Báo cáo tổng hợp theo kỳ ──
    with tab3:
        today = datetime.date.today()
        tu, den, td = chon_ky("ky", today)

        xl_chon = st.selectbox("Đề xuất mức xếp loại", XL_MUC, key="ky_xl")
        nhan_xet = st.text_area("Nhận xét, đánh giá bổ sung (không bắt buộc)", height=70)
        phuong_huong = st.text_area("Phương hướng nhiệm vụ tiếp theo (không bắt buộc)", height=70)

        d = tinh_toan_ky(tu, den, xl_chon)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Tổng công việc", d["tong"])
        c2.metric("Hoàn thành", f"{d['ht']} ({d['pt']}%)")
        c3.metric("Trễ hạn", d["tre"])
        c4.metric("Đề xuất xếp loại", d["xep_loai"])

        with st.expander("👁 Xem chi tiết công việc trong kỳ"):
            if d["tasks"]:
                df = pd.DataFrame([{"Ngày": _fmt_d(t["ngay_bd"]), "Nội dung": t["title"],
                                    "Loại": t.get("loai", ""), "Trạng thái": t["trang_thai"],
                                    "Kết quả": t.get("ket_qua") or ""} for t in d["tasks"]])
                st.dataframe(df, use_container_width=True, hide_index=True)
            else:
                st.info("Không có công việc trong kỳ này")

        if st.button("⬇ Xuất báo cáo Word theo kỳ", type="primary"):
            buf = build_ky_report(ten, chuc_vu, tu, den, td, d, nhan_xet, phuong_huong)
            st.download_button("📥 Tải file .docx",
                                data=buf, file_name=f"BaoCao_{td.replace(' ', '_').replace('/', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════════════════
#  TRANG: CÀI ĐẶT
# ══════════════════════════════════════════════════════
elif page == "⚙️ Cài đặt":
    st.title("⚙️ Cài đặt thông tin người dùng")
    new_ten = st.text_input("Họ và tên (*)", value=ten)
    new_cv = st.text_input("Chức vụ (hiển thị ở thanh bên)", value=chuc_vu)
    new_cv_dd = st.text_area(
        "Chức vụ, đơn vị đầy đủ (dùng trong Biểu 01 — mục \"Chức vụ, đơn vị\")",
        value=chuc_vu_day_du, height=70,
        placeholder="VD: Chuyên viên Văn phòng Ban Tuyên giáo và Dân vận Tỉnh ủy Tuyên Quang.")
    if st.button("💾 Lưu thay đổi", type="primary"):
        if not new_ten.strip():
            st.warning("Họ tên không được trống!")
        else:
            cfg_set("ten", new_ten.strip())
            cfg_set("chuc_vu", new_cv.strip() or "Chuyên viên")
            cfg_set("chuc_vu_day_du", new_cv_dd.strip())
            st.success("Đã cập nhật!")
            st.rerun()
